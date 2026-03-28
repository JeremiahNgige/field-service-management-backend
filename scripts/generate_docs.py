from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_document():
    doc = Document()

    # --- Document Styles Customization ---
    styles = doc.styles
    title_style = styles["Title"]
    title_font = title_style.font
    title_font.name = "Arial"
    title_font.size = Pt(28)
    title_font.color.rgb = RGBColor(0, 51, 102)

    heading1_style = styles["Heading 1"]
    h1_font = heading1_style.font
    h1_font.name = "Arial"
    h1_font.size = Pt(20)
    h1_font.color.rgb = RGBColor(0, 76, 153)

    heading2_style = styles["Heading 2"]
    h2_font = heading2_style.font
    h2_font.name = "Arial"
    h2_font.size = Pt(16)
    h2_font.color.rgb = RGBColor(0, 102, 204)

    # --- TITLE PAGE ---
    title = doc.add_heading(
        "Field Service Management (EZE)\nTechnical Documentation", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n\n\n\n")

    subtitle = doc.add_paragraph("Comprehensive Overview of Backend & Frontend Systems")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 1. EXECUTIVE SUMMARY ---
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "The Field Service Management (FSM) platform is a robust, end-to-end solution designed to handle "
        "the lifecycle of field operations. It facilitates job creation, dynamic technician dispatching, "
        "and real-time status updates."
    )
    doc.add_paragraph(
        "The system is composed of two primary architectures:\n"
        "1. A scalable Django REST framework backend orchestrating core business logic, asynchronous notifications, and S3 media storage.\n"
        "2. A premium Flutter mobile application featuring real-time state management, biometric security, and dynamic map routing."
    )

    # --- 2. SYSTEM ARCHITECTURE ---
    doc.add_heading("2. System Architecture & Data Flow", level=1)
    doc.add_paragraph(
        "The architecture follows a decoupled client-server model communicating exclusively via REST APIs secured by JSON Web Tokens (JWT)."
    )
    p = doc.add_paragraph()
    p.add_run("Key Flow Mechanics:\n").bold = True
    doc.add_paragraph(
        "• Client to Server: The Flutter app uses asynchronous generic repositories to communicate with the Django API, passing Auth tokens securely retrieved from encrypted local storage.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "• Event-Driven Actions: When a job is assigned to a technician, the database triggers a Django signal. This signal is handed off to a Celery Worker (using Redis as the message broker) to asynchronously dispatch a Firebase Cloud Messaging (FCM) push notification without blocking the HTTP response thread.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "• Media Storage (S3/MinIO): Heavy payloads (photos, signatures) are not routed through Django. Instead, the backend generates cryptographic S3 Presigned URLs, empowering the frontend to upload and download large binaries directly to the storage bucket.",
        style="List Bullet",
    )

    # --- 3. BACKEND INFRASTRUCTURE ---
    doc.add_heading("3. Backend Infrastructure (Django/Docker)", level=1)

    doc.add_heading("Technology Stack", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Component"
    hdr_cells[1].text = "Technology"

    tech_stack = [
        ("Framework", "Django 4.2 + Django REST Framework"),
        ("Database", "PostgreSQL 15"),
        ("Auth", "JWT (djangorestframework-simplejwt)"),
        ("Task Queue", "Celery + Redis"),
        ("Object Storage", "MinIO (S3-compatible)"),
        ("Push Notifications", "Firebase Admin SDK (FCM)"),
        ("Containerization", "Docker & Docker Compose"),
    ]
    for comp, tech in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = tech

    doc.add_heading("Core Models", level=2)
    p = doc.add_paragraph()
    p.add_run("User Model: ").bold = True
    p.add_run(
        "Role-based accounts supporting admin, technician, driver, and inspectors. Utilizes UUID primary keys and tracks FCM device tokens."
    )
    p = doc.add_paragraph()
    p.add_run("Job Model: ").bold = True
    p.add_run(
        "Tracks job lifecycle statuses (unassigned, assigned, in_progress, completed, cancelled), links to geospatial locations, schedules (start/end times), pricing, and media (signatures/photos)."
    )

    doc.add_heading("Testing & CI", level=2)
    doc.add_paragraph(
        "Extensive API Integration tests leverage APITestCase to validate end-to-end critical flows (auth workflows, job lifecycles, and view isolation constraints)."
    )

    # --- FUTURE CONSIDERATIONS (CI/CD) ---
    p_ci = doc.add_paragraph()
    p_ci.add_run("Future Considerations (CI/CD): ").bold = True
    p_ci.add_run(
        "Given more time, a robust Continuous Integration / Continuous Deployment (CI/CD) pipeline would be implemented "
        "(e.g., using GitHub Actions or GitLab CI) to automatically orchestrate and run the integration test suites across "
        "both the frontend Flutter repository and the Django backend before any code merges or deployment to production environments."
    )

    # --- 4. FRONTEND ARCHITECTURE ---
    doc.add_page_break()
    doc.add_heading("4. Frontend Architecture (Flutter)", level=1)

    doc.add_heading("UI/UX Philosophy", level=2)
    doc.add_paragraph(
        "The mobile application leverages modern premium aesthetics including vibrant gradient palettes, "
        "responsive glassmorphism overlays, and subtle micro-animations to enhance the user experience."
    )

    doc.add_heading("State Management & Networking", level=2)
    doc.add_paragraph(
        "• BLoC (Business Logic Component): Ensures strict separation of UI presentation code from underlying data mapping.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "• Dependency Injection: GetIt acts as the service locator providing singleton access to repositories and secure storage adapters.",
        style="List Bullet",
    )

    doc.add_heading("Native Security Integrations", level=2)
    doc.add_paragraph(
        "• Biometric Authentication: Incorporates platform-aware FaceID/TouchID integrations verifying users sequentially after initial JWT acquisition.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "• Secure Storage: Uses iOS Keychain and Android EncryptedSharedPreferences to house sensitive JWTs.",
        style="List Bullet",
    )

    doc.add_heading("Geolocator & Mapping", level=2)
    doc.add_paragraph(
        "• The app seamlessly calculates dynamic distance values utilizing the Geolocator package relative to targeted job coordinates.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "• Displays interactive embed maps tracking physical markers with native OS deeplinking to launch turn-by-turn navigation apps.",
        style="List Bullet",
    )

    # --- 5. DEPLOYMENT & DEVOPS ---
    doc.add_heading("5. Deployment (DevOps)", level=1)
    doc.add_paragraph(
        "The entire infrastructure is unified via a Docker Compose manifest. This ensures instantaneous environment parity "
        "across development, testing, and production phases."
    )

    doc.add_heading("Service Layout", level=2)
    doc.add_paragraph("The orchestrator manages 5 distinct containers:")
    doc.add_paragraph(
        "1. web: The Django Gunicorn WSGI server handling HTTP REST traffic."
    )
    doc.add_paragraph("2. db: The persistent PostgreSQL database container.")
    doc.add_paragraph("3. redis: In-memory datastore broker caching Celery signals.")
    doc.add_paragraph(
        "4. celery: Dedicated worker container processing background push notifications."
    )
    doc.add_paragraph("5. minio: S3 interface capturing directly-uploaded assets.")

    doc.add_paragraph("\n\n--- End of Documentation ---")
    p_end = doc.paragraphs[-1]
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


if __name__ == "__main__":
    out_file = "Field_Service_Management_Documentation.docx"
    doc = create_document()
    doc.save(out_file)
    print(f"Successfully generated {out_file}")
